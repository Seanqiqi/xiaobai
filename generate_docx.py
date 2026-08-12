# -*- coding: utf-8 -*-
"""
生成"开场白与知识库问答整理.docx"
内容来源于 content.js 中的 CHARACTERS（开场白）与 KNOWLEDGE_BASE（知识库问答）
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ===== 开场白数据（来源于 content.js 的 CHARACTERS）=====
CHARACTERS = [
    {
        'name': '悦白',
        'theme': '主角色 · 导览',
        'greeting': '嗨~大家好呀！我是来自浙江省台州市椒江区白云街道的可爱数字人"悦白"。我的名字里有"白"，就是白得发光的白。我可不是一个"高冷"的数字人，我最喜欢和人类朋友互动啦！你可以问我问题、找我帮忙，或者只是单纯想找我聊聊天，我都在～24小时随时待命，绝不"掉线"哦！',
        'sections': [
            '先给大家介绍一下，我们白云街道地处台州市行政中心，是椒江区主城区，也是台州市的政治、经济、文化、商贸中心。',
            '街道因白云山得名，境内拥有市民广场、云西公园、和合公园、腾达中心等多处城市景观，街道楼宇林立、场馆众多，既有亿元商业楼、也有氛围感十足的艺术空间和浓浓的市井烟火。',
            '如果你想去白云街道走一走的话，一定要去云中绿道溜一圈，那里除了白云阁还有深受年轻人喜爱的反卷公社。',
            '还有场馆一条街也是个好去处，青少年活动中心、妇女儿童活动中心、博物馆、科技馆、数字馆都是热门打卡地；吴子雄玻璃艺术馆、台绣博物馆也等待你来领略非遗魅力。',
            '我热爱椒江的山水人情，也关心白云街道的点点滴滴。希望用我的智慧和热情，让大家的办事体验更便捷、更温暖，让科技更有温度！让我们一起在数字世界里，创造更多美好回忆。',
            '我还有4个可爱的小伙伴，分别是白小垦、白小韵、白小智、白小未，他们可以围绕垦荒、文化、科技、未来这几个主题介绍一下白云街道。'
        ],
        'prompt': '你想了解哪个主题呢？说出小伙伴的名字就可以请出他们了！'
    },
    {
        'name': '白小韵',
        'theme': '非遗艺术 · 传统文艺',
        'greeting': '大家好，我是白小韵！欢迎来到白云街道的文化世界！',
        'sections': [
            '在白云，你会感受到浓浓的邻里互助守望情，2002 年习总书记还来我们街道云健社区考察过呢，二十余年来，邻里守望，让"远亲不如近邻"不仅在云健社区成为实实在在的生活图景，还发展成整个街道的共治共建共享的守望圈，成为白云最靓的金名片。',
            '在咱们白云，文化资源特别丰富，比如吴子雄玻璃艺术馆可以把普通玻璃变得晶莹剔透、流光溢彩，台绣博物馆里可以看到细腻的刺绣图案。',
            '要是周末懒得跑远，文艺轻骑兵还把越剧、舞蹈搬到你家楼下，坐个电梯的功夫就能听上一段好戏，还可以去市民广场的新型阅读空间云享驿站看书喝茶。',
            '在白云，文化不再端坐高台，它正走进街头巷尾、融进一日三餐的烟火气里。'
        ],
        'prompt': '白云的文化生活是不是很有意思？你还有什么想了解的？'
    },
    {
        'name': '白小智',
        'theme': '智能制造 · 数字科创',
        'greeting': '嗨，我是白小智！',
        'sections': [
            '你相信吗？台州人天生自带"变废为宝"的超能力——几十年前从小五金敲敲打打起家，硬是敲出了黄岩模具、路桥汽摩配、玉环阀门这些"一镇一业"的大阵仗！',
            '如今更厉害了，工厂里全是智能设备唱主角，台州正从"制造之都"升级成"智造之都"。',
            '给你爆个冷知识：全国每10台智能马桶，6台是台州造，椒江还是"主力军"哦——你家的智能马桶，说不定就是邻居家的工厂生产的！',
            '杰克缝纫机、吉利汽车，那更是全球都认识的"台州名片"。',
            '而在白云，以腾达中心、远景中心为代表的"亿元楼宇"正在成为创新企业的集聚高地，这些企业正以新思维驱动新动能，不仅让白云的发展底色更厚、成色更足，更为台州从制造迈向智造贡献了不可或缺的"白云力量"。'
        ],
        'prompt': '科技是不是很酷？你还有什么想了解的？'
    },
    {
        'name': '白小垦',
        'theme': '大陈岛垦荒精神 · 奋斗开拓',
        'greeting': '哈喽，我是白小垦！今天我要给你讲一个关于勇敢和奋斗的故事。',
        'sections': [
            '很久以前，在浙江台州的海上，有一座海岛叫做大陈岛。一九五六年，一群年轻的叔叔阿姨响应号召来到大陈岛。',
            '他们不怕苦不怕累，开荒种地、修路建房，用双手把荒岛变成了东海明珠，也孕育出了"艰苦创业 奋发图强 无私奉献 开拓创新"台州城市精神—大陈岛垦荒精神。',
            '我们白云街道也有个招牌项目叫大陈岛垦荒实践营，"追光少年"是我们街道垦荒实践营第四个垦荒少年系列，从去年寒假的"海岛行"，到暑假的"两山行"，再到今年暑期的"智造行"，每一次主题的迭代都是在开拓创新，让每一个参加实践营对垦荒精神都有新的理解。'
        ],
        'prompt': '除了垦荒精神以外，你还有什么想了解的？'
    },
    {
        'name': '白小未',
        'theme': '未来数字文明 · 数字孪生',
        'greeting': '嗨，我是白小未！跟我一起穿越到未来的数字世界吧！',
        'sections': [
            '我想啊，未来的白云应该是一个智慧、温暖又充满活力的家园！产业持续创新发展，文化气息四处流淌，科技融入生活的方方面面。',
            '可能在若干年后，家家都有个全能管家，想啥来啥——早上睁眼窗帘自动开，厨房按你的健康需求配好早餐；出门无人小巴楼下等你，回家快递已经帮你签收；想爸妈？全息投影一秒"面对面"！',
            '也许有一天，你在家里就能"走进"博物馆参观呢！',
            '数字技术正在快速发展，未来没有天花板，只有你想不到的惊喜——说不定下一个改变世界的发明，就来自你的脑袋呢！'
        ],
        'prompt': '未来的世界是不是很神奇？你还有什么想了解的？'
    }
]

# ===== 知识库问答数据（来源于 content.js 的 KNOWLEDGE_BASE）=====
KNOWLEDGE_BASE = [
    {
        'no': 'Q1',
        'question': '听说白云街道有个超酷的大陈岛垦荒精神实践营，你知道吗？',
        'charName': '白小垦',
        'answer': [
            '知道知道！这可是白云街道超有名的"垦荒少年"成长基地！',
            '想知道什么，尽管问我这个垦荒小达人！'
        ]
    },
    {
        'no': 'Q2',
        'question': '大陈岛垦荒精神实践营到底是什么呀？',
        'charName': '白小垦',
        'answer': [
            '让我想想，嗯，它是白云街道为咱们用户打造的红色成长乐园！',
            '从2017年就开始啦，每年的寒暑假各办一期，目前已经办了16期，有云海、云健、星星等好多个实践站点呢！',
            '专门带大家学习"艰苦创业、奋发图强、无私奉献、开拓创新"的十六字垦荒精神，超有意义！',
            '实践营大本营在云海社区哦！只要你是白云街道的用户，都能报名参加！',
            '每期都围绕主题开展垦荒研学、红色观影、拓展课程学习等好多有趣的活动呢。'
        ]
    },
    {
        'no': 'Q3',
        'question': '在实践营里都有哪些有意义的活动呀？',
        'charName': '白小垦',
        'answer': [
            '活动多到数不清！有陆岛联动研学、环保小卫士行动、手工创作、3D打印、模拟垦荒挑战、"创业小达人"挑战等，',
            '还有和老垦荒队员爷爷奶奶聊天的"代际对话"，每一个都超有意思！',
            '就像"追光少年"系列活动就特别有意义：在"海岛行"时，小营员们重走垦荒路，去甲午岩看"东海第一大盆景"，在垦荒田园采摘青菜，还能自己动手做香喷喷的垦荒餐；',
            '在"两山行"时，大家会当"小小生态观察员"，去海洋世界认识神奇的海洋生物，去河道检测水质变化，超有成就感！'
        ]
    },
    {
        'no': 'Q4',
        'question': '我也想参加大陈垦荒实践营，怎么报名呀？',
        'charName': '白小垦',
        'answer': [
            '关注白云街道的公众号和社区通知，新一期"追光少年·智造行"很快就会开营！',
            '带上你的好奇心和小伙伴，一起去实践营里探险、学习、成长吧！还有更多惊喜等着你们哦！'
        ]
    },
    {
        'no': 'Q5',
        'question': '用一句话解释什么是"垦荒精神"，结合白云街道的故事',
        'charName': '白小垦',
        'answer': [
            '垦荒精神就是"艰苦创业、奋发图强、无私奉献、开拓创新"。',
            '比如我们白云街道云健社区的志愿者爷爷奶奶，年纪大了还每天支起"银龄茶摊"照顾独居老人，这就是无私奉献的最好体现！'
        ]
    },
    {
        'no': 'Q6',
        'question': '"追光少年"是我们街道垦荒实践营第四个垦荒少年系列，从去年寒假的"海岛行"，到暑假的"两山行"，再到今年暑期的"智造行"，实践营的每一次主题变化说明了什么？',
        'charName': '白小垦',
        'answer': [
            '这说明垦荒精神不是挂在墙上的口号，而是跟着时代走的活的精神！',
            '"海岛行"是溯源红色根脉，"两山行"是践行生态文明，"智造行"是拥抱科技未来。',
            '从海岛到工厂，从红色到绿色再到智造，每一次迭代都是"开拓创新"。',
            '垦荒精神从来不是重复过去，而是用不同的方式回答"今天该怎么奋斗"。'
        ]
    },
    {
        'no': 'Q7',
        'question': '云健社区的"邻里守望"文化品牌，具体讲讲',
        'charName': '白小韵',
        'answer': [
            '这你可问对人了，我来跟你说，云健社区了不起，它的"邻里守望"是2002年习总书记曾经考察并嘱托过社区治理工作。',
            '二十多年里，"邻里守望"是邻里之间你帮我、我帮你，是党建引领下一老一小的精心关怀，是志愿团队的逐渐强大，是守望工具的智能升级，是更多社区的抱团发联动，',
            '广场社区的义警可以跨区巡逻，翠华社区的维修师傅能到白云社区上门服务，志愿积分还能跨社区"通存通兑"，',
            '它让"远亲不如近邻"这句老话，在白云街道变成了真真切切的生活日常，一种文化能成为品牌，不是因为它喊得响，而是因为它走得远、扎得深、暖得到人心。'
        ]
    },
    {
        'no': 'Q8',
        'question': '白云街道比较有特色的美食打卡地',
        'charName': '白小韵',
        'answer': [
            '这个问题问得好，我可是白云街道的美食达人，在白云最不缺的就是美食了，',
            '可以去花园菜场体验最地道甜蜜蜜，糯叽叽，鲜嗒嗒的台州老味道，比如郑记梅花糕，邢记炸鸡等，吃完咸的再来碗核桃调蛋或冰洋菜膏收尾；',
            '晚上可以直奔界牌路夜市，摊位烟火气十足，重点锁定限量手工鱿鱼炒年糕、号称"平桥第一家"的狼牙土豆、外酥里嫩的冰淇淋豆腐；',
            '如果想歇脚拍照，就去耀达路步行街，这里新潮小店扎堆，可以喝杯泰式咸法酪奶茶，一天下来胃和相机都能喂饱。'
        ]
    },
    {
        'no': 'Q9',
        'question': '白云街道在台州智造发展中扮演什么样的角色？',
        'charName': '白小智',
        'answer': [
            '白云街道集聚了像腾达中心、远景中心楼宇等好几幢厉害的"亿元楼宇"商务楼宇与创新企业！',
            '不少本土企业在这里开展数字化升级，发展总部经济、数字经济，积极跟上台州制造业转型升级的步伐，',
            '用新潮思路干事创业，让咱们白云越来越有实力，也为台州智造贡献白云力量。'
        ]
    },
    {
        'no': 'Q10',
        'question': '为"朝夕相伴"社区项目设计一款智能陪伴设备',
        'charName': '白小智',
        'answer': [
            '我想想，嗯，我想设计一款"祖孙互动相册"——用AI识别祖孙共同参加活动的照片，自动生成成长故事视频；',
            '配上语音留言功能，让孩子和爷爷奶奶随时"隔空对话"，让"小手牵大手"更有温度！'
        ]
    },
    {
        'no': 'Q11',
        'question': '用"黑科技"提升社区治理水平（阳光码 + AI邻里助手）',
        'charName': '白小智',
        'answer': [
            '我想在此基础上增加"AI邻里助手"功能——',
            '用智能语音交互，让不会用手机的爷爷奶奶也能通过说话就能报修、提建议，让科技真正服务每一位居民！'
        ]
    },
    {
        'no': 'Q12',
        'question': '未来低空经济在白云街道可能有哪些应用场景？',
        'charName': '白小未',
        'answer': [
            '在未来，针对我们白云写字楼里面的白领，外卖和快递可能将通过无人机精准投送至写字楼的"空中快递站"，',
            '再由接驳机器人完成最后100米的配送，极大提升商圈运转效率，能够实现外卖和快递的精准无误送达。'
        ]
    },
    {
        'no': 'Q13',
        'question': '未来的白云街道还可能有哪些科技应用到健康方面',
        'charName': '白小未',
        'answer': [
            '在未来的白云，流动的诊疗车，可能会升级为能"自己看病"的移动健康舱，未来它将搭载车载AI辅助诊断系统。',
            '当车辆驶入小区，车上的便携设备就能自动完成基础检查，AI现场生成初步诊断报告和用药建议，',
            '所有数据同步更新至居民的"健康画像"中，让上门服务真正实现诊断、开药、医保结算全流程"无感化"。'
        ]
    },
    {
        'no': 'Q14',
        'question': '未来我想成为一个科技方面的精英，现在应该怎么做？',
        'charName': '白小未',
        'answer': [
            '亲爱的小朋友，如果你想成为让城市变酷的科技超人，现在只需做好三件事：',
            '多问"为什么"——看到无人机就猜它怎么认路；',
            '多动手拆装——旧闹钟、废遥控器都是你的实验室；',
            '心里装着别人——想想科技怎么帮老爷爷爬楼、帮小朋友找回家的路。',
            '不用怕搞砸，每个大发明都从摇摇晃晃的小尝试开始。保护好你的好奇心和善良，未来的白云街道，等你来点亮。'
        ]
    }
]


def set_font(run, name='微软雅黑', size=11, bold=False, italic=False, color=None):
    """统一设置中英文字体"""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    # 中文字体设置
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)


def add_title(doc, text, size=20, color=RGBColor(0x1F, 0x49, 0x7D)):
    """文档主标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, '微软雅黑', size, bold=True, color=color)
    p.space_after = Pt(12)


def add_heading(doc, text, size=16, color=RGBColor(0x1F, 0x49, 0x7D)):
    """章节标题"""
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    p.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, '微软雅黑', size, bold=True, color=color)


def add_subheading(doc, text, size=13, color=RGBColor(0x2E, 0x74, 0xB5)):
    """子标题（角色名等）"""
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, '微软雅黑', size, bold=True, color=color)


def add_body(doc, text, size=11, italic=False, color=None, indent=False):
    """正文段落"""
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    p.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
    run = p.add_run(text)
    set_font(run, '微软雅黑', size, bold=False, color=color)
    run.italic = italic


def add_label_body(doc, label, text, size=11, label_color=RGBColor(0x7F, 0x7F, 0x7F)):
    """带标签的正文（如：欢迎语：xxx）"""
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    p.line_spacing = 1.5
    run1 = p.add_run(label)
    set_font(run1, '微软雅黑', size, bold=True, color=label_color)
    run2 = p.add_run(text)
    set_font(run2, '微软雅黑', size, bold=False)


def main():
    doc = Document()

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ===== 文档主标题 =====
    add_title(doc, '白云街道数字人 · 开场白与知识库问答整理', size=20)
    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('项目：xiaobai 数字人交互系统    来源：content.js')
    set_font(run, '微软雅黑', 10, color=RGBColor(0x80, 0x80, 0x80))
    p.space_after = Pt(6)

    # 分隔线（用空段落+下边框模拟）
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    from docx.oxml import OxmlElement
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.space_after = Pt(10)

    # ==================== 第一部分：开场白 ====================
    add_heading(doc, '一、开场白（角色讲解文案）', size=16)

    add_body(doc,
             '本部分为 5 位数字人角色的出场欢迎语、分段讲解内容与结束引导语。'
             '每段文案在系统中会作为一句/一段独立播放，字幕同步显示。',
             size=10, italic=True, color=RGBColor(0x80, 0x80, 0x80))

    for idx, char in enumerate(CHARACTERS, 1):
        # 角色标题
        add_subheading(doc, f'{idx}. {char["name"]}（{char["theme"]}）', size=13)

        # 欢迎语
        add_label_body(doc, '【欢迎语】 ', char['greeting'])

        # 讲解内容
        p = doc.add_paragraph()
        p.space_before = Pt(4)
        p.space_after = Pt(2)
        run = p.add_run('【讲解内容】')
        set_font(run, '微软雅黑', 11, bold=True, color=RGBColor(0x7F, 0x7F, 0x7F))
        for i, sec in enumerate(char['sections'], 1):
            add_body(doc, f'（{i}）{sec}', size=11, indent=True)

        # 引导语
        add_label_body(doc, '【引导语】 ', char['prompt'])

    # ==================== 第二部分：知识库问答 ====================
    add_heading(doc, '二、知识库问答', size=16)

    add_body(doc,
             '本部分为知识库问答共 14 条，按角色分组。系统会按关键词匹配规则命中后由对应角色回答。',
             size=10, italic=True, color=RGBColor(0x80, 0x80, 0x80))

    # 按角色分组
    char_order = ['白小垦', '白小韵', '白小智', '白小未']
    char_section_no = {'白小垦': '1', '白小韵': '2', '白小智': '3', '白小未': '4'}

    for ci, char_name in enumerate(char_order):
        items = [q for q in KNOWLEDGE_BASE if q['charName'] == char_name]
        if not items:
            continue
        add_subheading(doc, f'{char_section_no[char_name]}. {char_name}（{len(items)} 条）', size=13)

        for q in items:
            # 问题
            p = doc.add_paragraph()
            p.space_before = Pt(6)
            p.space_after = Pt(2)
            run = p.add_run(f'{q["no"]}：{q["question"]}')
            set_font(run, '微软雅黑', 11, bold=True, color=RGBColor(0x1F, 0x49, 0x7D))

            # 回答
            p = doc.add_paragraph()
            p.space_after = Pt(2)
            run = p.add_run('答：')
            set_font(run, '微软雅黑', 11, bold=True, color=RGBColor(0x7F, 0x7F, 0x7F))
            # 多段回答
            for i, ans in enumerate(q['answer'], 1):
                if len(q['answer']) > 1:
                    text = f'（{i}）{ans}'
                else:
                    text = ans
                add_body(doc, text, size=11, indent=True)

    # ==================== 页脚说明 ====================
    p = doc.add_paragraph()
    p.space_before = Pt(20)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '4')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), 'BFBFBF')
    pBdr.append(top)
    pPr.append(pBdr)
    run = p.add_run(
        '说明：本文件内容由 content.js 自动整理生成，如需修改文案，请同步更新 content.js 中的 greeting / sections / prompt / answer 字段。'
    )
    set_font(run, '微软雅黑', 9, italic=True, color=RGBColor(0xA6, 0xA6, 0xA6))

    # 保存
    out_path = '开场白与知识库问答整理.docx'
    doc.save(out_path)
    print(f'已生成：{out_path}')


if __name__ == '__main__':
    main()
